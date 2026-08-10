"""
Generate a clean, NON-TECHNICAL Access Level Confirmation document.
Plain English only — suitable for client/user review and sign-off.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEST = "/Users/deeproot/data/21MARCH2026/my-ats-platform/docs/WVA_CRM_Access_Level_Confirmation.docx"

BRAND_ORANGE = RGBColor(0xE8, 0x8E, 0x2E)
BRAND_DARK   = RGBColor(0x1E, 0x29, 0x3B)
GREEN        = RGBColor(0x16, 0xA3, 0x4A)
RED          = RGBColor(0xDC, 0x26, 0x26)
BLUE         = RGBColor(0x1D, 0x4E, 0xD8)
GREY_TEXT    = RGBColor(0x64, 0x74, 0x8B)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_p(cell, text, size=10, bold=False, colour=None,
           align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    if colour: r.font.color.rgb = colour

def add_para(text="", size=10, bold=False, colour=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=5, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size)
        if colour: r.font.color.rgb = colour
    return p

def section_heading(text, colour=BRAND_DARK):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after  = Pt(5)
    for run in h.runs:
        run.font.color.rgb = colour
        run.font.size      = Pt(13)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
t1 = doc.add_heading("WorkVision Australia – CRM System", level=1)
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t1.runs:
    r.font.color.rgb = BRAND_DARK
    r.font.size = Pt(22)

t2 = doc.add_heading("Staff Access Levels – Please Review & Confirm", level=2)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t2.runs:
    r.font.color.rgb = BRAND_ORANGE
    r.font.size = Pt(14)

add_para("July 2026   |   Prepared for your review and approval",
         size=9, colour=GREY_TEXT, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
add_para("─" * 100, size=8, colour=GREY_TEXT,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

# ══════════════════════════════════════════════════════════════════════════════
# INTRO
# ══════════════════════════════════════════════════════════════════════════════
section_heading("About This Document")
add_para(
    "This document outlines who can access which areas of the WVA CRM system based on their role. "
    "We have set up two types of staff access — Staff and Training Admin — along with the existing Admin role. "
    "Please read through each section, answer the questions, and sign at the end so we can proceed.",
    size=10, after=10)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — THE TWO STAFF ROLES
# ══════════════════════════════════════════════════════════════════════════════
section_heading("1.  The Two Staff Roles")
add_para(
    "There are three roles in the CRM. Each role determines which parts of the system "
    "a staff member can see and use:",
    size=10, after=6)

roles = [
    ("🔑  Admin",
     "Full access to everything in the CRM, including the ability to "
     "add new staff, change access levels, and manage all records.",
     "F0FDF4", GREEN),
    ("👤  Staff\n(Consultants / Recruiters)",
     "Full access to all seven main areas of the CRM:\n"
     "Dashboard  ·  Important Updates  ·  Employers  ·  Vacancies  ·  "
     "Candidates  ·  Placements  ·  Providers",
     "EFF6FF", BLUE),
    ("📋  Training Admin",
     "Limited access — can only view the Candidates section and manage the Training section.\n"
     "All other areas of the CRM will be hidden from this role.",
     "FFF7ED", BRAND_ORANGE),
]

for title, desc, bg, colour in roles:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(4.5)
    tbl.columns[1].width = Cm(12.5)
    row = tbl.rows[0].cells
    set_cell_bg(row[0], bg)
    set_cell_bg(row[1], "FFFFFF")
    cell_p(row[0], title, size=10, bold=True, colour=colour, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_p(row[1], desc, size=10)
    add_para(after=4)

add_para(after=4)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — MODULE ACCESS TABLE
# ══════════════════════════════════════════════════════════════════════════════
section_heading("2.  What Each Role Can Access")
add_para("✅ Full Access     👁 View Only (cannot add or edit)     ❌ Not Available",
         size=9, colour=GREY_TEXT, italic=True, after=6)

modules = [
    ("Dashboard",         "✅ Full", "✅ Full", "❌ Not Available"),
    ("Important Updates", "✅ Full", "✅ Full", "❌ Not Available"),
    ("Employers",         "✅ Full", "✅ Full", "❌ Not Available"),
    ("Vacancies",         "✅ Full", "✅ Full", "❌ Not Available"),
    ("Candidates",        "✅ Full", "✅ Full", "👁 View Only"),
    ("Placements",        "✅ Full", "✅ Full", "❌ Not Available"),
    ("Providers",         "✅ Full", "✅ Full", "❌ Not Available"),
    ("Training",          "✅ Full", "✅ Full", "✅ Full"),
    ("Reports",           "✅ Full", "❌ Not Available", "❌ Not Available"),
    ("User Management",   "✅ Full", "❌ Not Available", "❌ Not Available"),
]

tbl2 = doc.add_table(rows=1, cols=4)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([5.0, 3.5, 3.5, 5.0]):
    tbl2.columns[i].width = Cm(w)

hdr2 = tbl2.rows[0].cells
for i, h in enumerate(["Area / Module", "Admin", "Staff", "Training Admin"]):
    set_cell_bg(hdr2[i], "1E293B")
    cell_p(hdr2[i], h, bold=True, colour=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

def acc_col(v):
    if "✅" in v: return GREEN
    if "👁" in v: return BLUE
    return RED

for i, (mod, adm, stf, tra) in enumerate(modules):
    row = tbl2.add_row().cells
    bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
    for c in row: set_cell_bg(c, bg)
    cell_p(row[0], mod, size=10, bold=True)
    cell_p(row[1], adm, size=10, colour=acc_col(adm), align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_p(row[2], stf, size=10, colour=acc_col(stf), align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_p(row[3], tra, size=10, colour=acc_col(tra), align=WD_ALIGN_PARAGRAPH.CENTER)

add_para(after=8)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — ADMIN CONTROL
# ══════════════════════════════════════════════════════════════════════════════
section_heading("3.  How Admin Manages Access Levels")
add_para(
    "The Admin can manage all staff access levels directly inside the CRM — "
    "no outside help is needed. Here is what Admin can do:",
    size=10, after=6)

admin_items = [
    ("See all staff",        "View a list of all staff members, their role, and whether their account is active."),
    ("Change a role",        "Click on any staff member and select a new role from a simple list. The change happens straight away."),
    ("Disable an account",   "If a staff member leaves, Admin can disable their account without losing any of their history."),
    ("Add new staff",        "Admin can create a new account for a new team member and assign their role straight away."),
]

tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.columns[0].width = Cm(4.5)
tbl3.columns[1].width = Cm(12.5)

hdr3 = tbl3.rows[0].cells
for i, h in enumerate(["Admin Can", "Description"]):
    set_cell_bg(hdr3[i], "1E293B")
    cell_p(hdr3[i], h, bold=True, colour=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (cap, desc) in enumerate(admin_items):
    row = tbl3.add_row().cells
    bg = "F0FDF4" if i % 2 == 0 else "FAFFFE"
    for c in row: set_cell_bg(c, bg)
    cell_p(row[0], cap, size=10, bold=True, colour=GREEN)
    cell_p(row[1], desc, size=10, colour=GREY_TEXT)

add_para(after=8)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — QUESTIONS FOR USER
# ══════════════════════════════════════════════════════════════════════════════
section_heading("4.  We Need Your Feedback — Please Answer Below", colour=BRAND_ORANGE)
add_para(
    "Before we set this up, we have a few quick questions. "
    "Please tick your preferred answer for each one:",
    size=10, after=6)

questions = [
    ("1", "Should the Training Admin be able to see a candidate's phone number and email address?",
     "☐  Yes, show them          ☐  No, keep them hidden"),
    ("2", "Should the Training Admin be able to search for candidates by suburb or provider?",
     "☐  Yes, full search        ☐  No, limited search only"),
    ("3", "Should Staff (Consultants) be able to delete records such as candidates and placements?",
     "☐  Yes, Staff can delete   ☐  No, only Admin can delete"),
    ("4", "Is the list of areas for Staff correct?\n(Dashboard, Important Updates, Employers, Vacancies, Candidates, Placements, Providers)",
     "☐  Yes, this is correct    ☐  No, I need to make changes"),
    ("5", "Is the Training Admin access correct?\n(Candidates – view only, and Training – full access)",
     "☐  Yes, this is correct    ☐  No, I need to make changes"),
    ("6", "Are there any other staff members who need a different type of access not covered above?",
     "☐  No, these roles cover everyone\n☐  Yes (please write details in the notes box below)"),
]

tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = "Table Grid"
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([0.8, 8.5, 7.7]):
    tbl4.columns[i].width = Cm(w)

hdr4 = tbl4.rows[0].cells
for i, h in enumerate(["#", "Question", "Your Answer"]):
    set_cell_bg(hdr4[i], "1E293B")
    cell_p(hdr4[i], h, bold=True, colour=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (num, qtext, opts) in enumerate(questions):
    row = tbl4.add_row().cells
    bg = "FFF8F1" if i % 2 == 0 else "FFFCF8"
    for c in row: set_cell_bg(c, bg)
    cell_p(row[0], num, size=10, bold=True, colour=BRAND_ORANGE,
           align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_p(row[1], qtext, size=10)
    cell_p(row[2], opts, size=10, colour=GREY_TEXT)

add_para(after=6)

# Notes box
add_para("Additional Notes or Requests:", size=10, bold=True, colour=BRAND_DARK, after=3)
tbl5 = doc.add_table(rows=1, cols=1)
tbl5.style = "Table Grid"
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl5.columns[0].width = Cm(17.0)
set_cell_bg(tbl5.rows[0].cells[0], "FFFDF5")
tbl5.rows[0].cells[0].paragraphs[0].add_run("\n\n\n\n\n\n").font.size = Pt(10)

add_para(after=10)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — SIGN-OFF
# ══════════════════════════════════════════════════════════════════════════════
section_heading("5.  Sign-Off")
add_para(
    "By signing below, you confirm that you have reviewed the access levels in this document "
    "and are happy for us to proceed with setting them up.",
    size=10, after=10)

tbl6 = doc.add_table(rows=2, cols=4)
tbl6.style = "Table Grid"
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([4.5, 4.5, 4.0, 4.0]):
    tbl6.columns[i].width = Cm(w)

hdr6 = tbl6.rows[0].cells
for i, h in enumerate(["Full Name", "Signature", "Date", "Position"]):
    set_cell_bg(hdr6[i], "1E293B")
    cell_p(hdr6[i], h, bold=True, colour=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

for c in tbl6.rows[1].cells:
    set_cell_bg(c, "FFFDF5")
    c.paragraphs[0].add_run("\n\n\n").font.size = Pt(10)

add_para(after=8)

fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    "WorkVision Australia CRM  ·  Staff Access Level Confirmation  ·  July 2026"
)
fr.font.size = Pt(8)
fr.font.color.rgb = GREY_TEXT
fr.italic = True

doc.save(DEST)
print(f"✅ Saved: {DEST}")
